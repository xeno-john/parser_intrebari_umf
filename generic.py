from docx import Document
from bs4 import BeautifulSoup
import re
import sys
import os

def parse_and_create_word(nume_examen, nume_fisier_html, nume_docx):
    if not nume_docx.endswith(".docx"):
        print("Numele documentului trebuie sa se termine in 'docx'!")
        return

    if not nume_fisier_html.endswith(".html"):
        print("Numele paginii trebuie sa se termine in 'html'!")
        return 

    if not os.path.exists(nume_fisier_html):
        print("Fisierul html trebuie sa fie in acelasi folder cu scriptul Python!")
        return

    document = Document()
    document.add_heading(nume_examen, 0)

    d = {"class": "mb-4"}

    with open(nume_fisier_html, "r", encoding="utf8") as f:
        soup = BeautifulSoup(f.read())
        html = soup.find(id="questionsSection").div.ol.find_next("li", **d)

        while True:
            intrebare = re.sub(r'\s+', ' ', html.div.text)
            titlu_intrebare = document.add_paragraph(style='List Number')
            runner = titlu_intrebare.add_run(intrebare)
            runner.bold = True

            for raspuns in html.ol.findAll('li'):
                eliminat_spatii = re.sub(r'\s+', ' ', raspuns.text)
                document.add_paragraph(
                    eliminat_spatii, style='List Bullet'
                )

            html = html.find_next("li", **d)

            if not html:
                break

    document.save(nume_docx)


if __name__ == "__main__":
    if not len(sys.argv) == 4:
        print("Numar incorect de argumente! Folositi <python generic.py nume_examen pagina_descarcata.html nume_word.docx>")
    else:
        parse_and_create_word(sys.argv[1], sys.argv[2], sys.argv[3])
